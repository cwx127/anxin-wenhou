from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
SOURCE = DOCS / "项目说明文档-内容稿.md"
OUTPUT = DOCS / "安心问候-项目说明文档.docx"
SCREENSHOTS = DOCS / "screenshots"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "18201D"
MUTED = "5D6963"
GREEN = "176B46"
GREEN_SOFT = "E8F4ED"
LINE = "D8E0DC"
CALLOUT = "F4F6F9"
AMBER_SOFT = "FFF3DC"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LINE, size="6", val="single"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), val)
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths, indent=120):
    total = sum(widths)
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def create_decimal_numbering(doc, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if bullet else "%1.")
    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "280")
    p_pr.extend([tabs, indent])
    level.extend([start, num_fmt, level_text, level_justification, p_pr])
    if bullet:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Arial")
        r_fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(r_fonts)
        level.append(r_pr)
    abstract.append(level)
    first_num_index = next(
        (position for position, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text_node, fld_end])
    set_run_font(run, size=9, color=MUTED)


def set_paragraph_border(paragraph, color=GREEN, size="18", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)


def clean_inline(text):
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    text = text.replace("`", "")
    return text


def add_inline_runs(paragraph, text, size=11, color=INK):
    text = clean_inline(text)
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, color=color)
        raw = match.group(0)
        bold = raw.startswith("**")
        value = raw[2:-2] if bold else raw[1:-1]
        run = paragraph.add_run(value)
        set_run_font(run, size=size, color=color, bold=bold, italic=not bold)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color)


def add_body_paragraph(doc, text, style="Normal", before=None, after=None):
    paragraph = doc.add_paragraph(style=style)
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(before)
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
    add_inline_runs(paragraph, text)
    return paragraph


def add_callout(doc, text, fill=CALLOUT, accent=GREEN, label=None):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=fill, size="2")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    set_paragraph_border(paragraph, color=accent, size="20", space="10")
    if label:
        run = paragraph.add_run(label + "  ")
        set_run_font(run, size=10, color=accent, bold=True)
    add_inline_runs(paragraph, text, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, rows):
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    widths_by_cols = {
        2: [2700, 6660],
        3: [1800, 2520, 5040],
        4: [1500, 2100, 2760, 3000],
        5: [1320, 1900, 2100, 1920, 2120],
    }
    widths = widths_by_cols.get(cols, [9360 // cols] * cols)
    if len(widths) == cols:
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths, indent=120)
    for r_index, row_values in enumerate(rows):
        for c_index, value in enumerate(row_values):
            cell = table.cell(r_index, c_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            add_inline_runs(paragraph, clean_inline(value), size=9.3 if r_index else 9.5, color=INK)
            if r_index == 0:
                set_cell_shading(cell, "F4F6F9")
                for run in paragraph.runs:
                    run.bold = True
            set_cell_border(cell, color=LINE, size="4")
        if r_index == 0:
            mark_header_row(table.rows[r_index])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F4F7")
    set_cell_border(cell, color="D8E0DC", size="4")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=9, color="344054")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image_with_caption(doc, path, caption):
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(6.0))
    inline_shape._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(10)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=9, color=MUTED, italic=True)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("安心问候  |  60 秒安心回执")
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    run = footer.add_run("初赛项目说明  ·  ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(footer)


def add_cover(doc):
    for _ in range(2):
        doc.add_paragraph().paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("阿里云“小有可为”养老方向  |  初赛材料")
    set_run_font(run, size=11, color=GREEN, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("安心问候")
    set_run_font(run, size=30, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("60 秒安心回执")
    set_run_font(run, size=18, color=GREEN, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("老人每天说一句或点一下，家属收到一张带原话证据的状态回执。")
    set_run_font(run, size=12, color=MUTED)

    table = doc.add_table(rows=2, cols=4)
    set_table_geometry(table, [1450, 3230, 1450, 3230], indent=120)
    values = [
        ("材料类型", "项目说明文档", "版本", "前端演示 MVP"),
        ("验收日期", "2026-08-08", "核心承诺", "不监控 · 不诊断 · 只同步必要信息"),
    ]
    for r, row in enumerate(values):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = ""
            if c in (0, 2):
                set_cell_shading(cell, GREEN_SOFT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, size=9.5, color=GREEN if c in (0, 2) else INK, bold=c in (0, 2))

    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    add_callout(doc, "本项目把“安心”定义为一条可核验的照护信息链：原话 → 事实 → 状态 → 人工动作。系统不把未回应写成危险，也不把规则回执写成诊断。", fill=GREEN_SOFT, accent=GREEN, label="核心判断")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(56)
    run = p.add_run("项目说明文档 · 提交前版本")
    set_run_font(run, size=10, color=MUTED)
    doc.add_page_break()


def parse_table_row(line):
    stripped = line.strip().strip("|")
    return [item.strip() for item in stripped.split("|")]


def is_separator(line):
    return bool(re.match(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?$", line.strip()))


def build_body(doc, lines):
    index = 0
    in_code = False
    code_lines = []
    previous_type = None
    active_number_id = None
    active_bullet_id = None
    while index < len(lines):
        raw = lines[index].rstrip("\n")
        stripped = raw.strip()
        if not stripped:
            previous_type = None
            index += 1
            continue
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if stripped.startswith("# "):
            index += 1
            continue
        if not any(line.strip().startswith("## 1.") for line in lines[: index + 1]):
            index += 1
            continue
        if stripped.startswith(">"):
            text = stripped.lstrip("> ")
            add_callout(doc, text, fill=CALLOUT, accent=GREEN)
            previous_type = "callout"
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = [parse_table_row(line) for line in table_lines if not is_separator(line)]
            if rows and len(rows[0]) >= 2:
                width = len(rows[0])
                rows = [row + [""] * (width - len(row)) for row in rows]
                rows = [row[:width] for row in rows]
                add_table(doc, rows)
            previous_type = "table"
            continue
        match = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if match:
            level = len(match.group(1)) - 1
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            paragraph.paragraph_format.keep_with_next = True
            add_inline_runs(paragraph, match.group(2), size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE)
            previous_type = "heading"
            index += 1
            continue
        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        number = re.match(r"^\d+\.\s+(.*)$", stripped)
        if bullet or number:
            if bullet:
                if previous_type != "bullet":
                    active_bullet_id = create_decimal_numbering(doc, bullet=True)
                paragraph = doc.add_paragraph(style="Normal")
                apply_numbering(paragraph, active_bullet_id)
                paragraph.paragraph_format.left_indent = Inches(0.375)
                paragraph.paragraph_format.first_line_indent = Inches(-0.194)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.208
            else:
                if previous_type != "number":
                    active_number_id = create_decimal_numbering(doc)
                paragraph = doc.add_paragraph(style="Normal")
                apply_numbering(paragraph, active_number_id)
                paragraph.paragraph_format.left_indent = Inches(0.375)
                paragraph.paragraph_format.first_line_indent = Inches(-0.194)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.208
            paragraph.paragraph_format.keep_together = True
            add_inline_runs(paragraph, (bullet or number).group(1))
            previous_type = "bullet" if bullet else "number"
            index += 1
            continue
        paragraph = doc.add_paragraph(style="Normal")
        add_inline_runs(paragraph, stripped)
        previous_type = "paragraph"
        index += 1


def add_appendix(doc):
    h = doc.add_paragraph(style="Heading 1")
    add_inline_runs(h, "附录 A：最终实测截图", size=16, color=BLUE)
    add_body_paragraph(doc, "以下截图来自 2026-08-08 的最终浏览器验收：老人端视口为 1440×1200，家属端视口为 1440×1000。截图只用于展示当前前端 MVP 的真实状态，不代表线上服务已部署。")
    add_image_with_caption(doc, SCREENSHOTS / "senior-dashboard.png", "图 A-1  老人端：今日问安入口、语音/文字降级与三个大按钮")
    add_image_with_caption(doc, SCREENSHOTS / "family-dashboard.png", "图 A-2  家属端：需关注回执、老人原话与可追溯依据")


def main():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_cover(doc)
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    build_body(doc, source_lines)
    add_appendix(doc)
    doc.core_properties.title = "安心问候：项目说明文档"
    doc.core_properties.subject = "60 秒安心回执前端演示 MVP"
    doc.core_properties.author = "安心问候项目组"
    doc.core_properties.keywords = "养老,安心问候,安心回执,MVP"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
